"""
Generate a clean, ATS-friendly resume as a downloadable .docx and .pdf
from structured form data (see modules/resume_builder.py for the schema).

No dependency on Microsoft Word / LibreOffice being installed:
  * .docx is built with python-docx
  * .pdf  is built with reportlab
so this works identically on any OS, including a bare department PC.
"""
from __future__ import annotations

import uuid
from typing import Any

import config


def _safe(text: Any) -> str:
    return str(text).strip() if text else ""


def _link(text: str, href: str) -> str:
    """Wrap text in a reportlab hyperlink markup tag, so email/LinkedIn/GitHub
    are clickable in the exported PDF instead of dead text."""
    return f'<link href="{href}"><u>{text}</u></link>'


def _contact_markup(data: dict) -> str:
    parts = []
    email = _safe(data.get("email"))
    if email:
        parts.append(_link(email, f"mailto:{email}"))
    phone = _safe(data.get("phone"))
    if phone:
        parts.append(phone)
    location = _safe(data.get("location"))
    if location:
        parts.append(location)
    linkedin = _safe(data.get("linkedin"))
    if linkedin:
        href = linkedin if linkedin.startswith("http") else f"https://{linkedin}"
        parts.append(_link(linkedin, href))
    github = _safe(data.get("github"))
    if github:
        href = github if github.startswith("http") else f"https://{github}"
        parts.append(_link(github, href))
    return " | ".join(parts)


def build_docx(data: dict) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(_safe(data.get("full_name")) or "Your Name")
    run.bold = True
    run.font.size = Pt(20)

    contact_bits = [
        _safe(data.get("email")),
        _safe(data.get("phone")),
        _safe(data.get("location")),
        _safe(data.get("linkedin")),
        _safe(data.get("github")),
    ]
    contact = doc.add_paragraph(" | ".join(b for b in contact_bits if b))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if data.get("summary"):
        _add_heading(doc, "Professional Summary")
        doc.add_paragraph(_safe(data["summary"]))

    if data.get("skills"):
        _add_heading(doc, "Skills")
        doc.add_paragraph(", ".join(s.strip() for s in data["skills"] if s.strip()))

    if data.get("experience"):
        _add_heading(doc, "Experience")
        for exp in data["experience"]:
            p = doc.add_paragraph()
            p.add_run(f"{_safe(exp.get('role'))} — {_safe(exp.get('company'))}").bold = True
            p.add_run(f"   ({_safe(exp.get('duration'))})")
            for bullet in exp.get("bullets", []):
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style="List Bullet")

    if data.get("projects"):
        _add_heading(doc, "Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            p.add_run(_safe(proj.get("title"))).bold = True
            if proj.get("tech"):
                p.add_run(f"   [{_safe(proj.get('tech'))}]")
            for bullet in proj.get("bullets", []):
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style="List Bullet")

    if data.get("education"):
        _add_heading(doc, "Education")
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.add_run(f"{_safe(edu.get('degree'))} — {_safe(edu.get('institution'))}").bold = True
            p.add_run(f"   ({_safe(edu.get('duration'))})")
            if edu.get("score"):
                doc.add_paragraph(f"Score: {_safe(edu.get('score'))}")

    if data.get("certifications"):
        _add_heading(doc, "Certifications")
        for c in data["certifications"]:
            if c.strip():
                doc.add_paragraph(c.strip(), style="List Bullet")

    out_path = config.GENERATED_DIR / f"resume_{uuid.uuid4().hex}.docx"
    doc.save(out_path)
    return str(out_path)


def _add_heading(doc, text: str) -> None:
    from docx.shared import Pt, RGBColor

    h = doc.add_paragraph()
    run = h.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(2)


def build_pdf(data: dict) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    out_path = config.GENERATED_DIR / f"resume_{uuid.uuid4().hex}.pdf"
    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm, topMargin=15 * mm, bottomMargin=15 * mm,
    )
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        "Section", parent=styles["Heading2"], textColor=HexColor("#1F4E79"), spaceBefore=10, spaceAfter=4
    )
    name_style = ParagraphStyle("Name", parent=styles["Title"], spaceAfter=2)
    body = styles["BodyText"]

    contact_style = ParagraphStyle("Contact", parent=body, textColor=HexColor("#1F4E79"))
    story = [
        Paragraph(_safe(data.get("full_name")) or "Your Name", name_style),
        Paragraph(_contact_markup(data), contact_style),
        Spacer(1, 6),
    ]

    if data.get("summary"):
        story += [Paragraph("PROFESSIONAL SUMMARY", heading_style), Paragraph(_safe(data["summary"]), body)]

    if data.get("skills"):
        story += [Paragraph("SKILLS", heading_style), Paragraph(", ".join(s.strip() for s in data["skills"] if s.strip()), body)]

    if data.get("experience"):
        story.append(Paragraph("EXPERIENCE", heading_style))
        for exp in data["experience"]:
            story.append(Paragraph(f"<b>{_safe(exp.get('role'))} — {_safe(exp.get('company'))}</b> ({_safe(exp.get('duration'))})", body))
            for bullet in exp.get("bullets", []):
                if bullet.strip():
                    story.append(Paragraph(f"• {bullet.strip()}", body))

    if data.get("projects"):
        story.append(Paragraph("PROJECTS", heading_style))
        for proj in data["projects"]:
            title = _safe(proj.get("title"))
            tech = f" [{_safe(proj.get('tech'))}]" if proj.get("tech") else ""
            story.append(Paragraph(f"<b>{title}</b>{tech}", body))
            for bullet in proj.get("bullets", []):
                if bullet.strip():
                    story.append(Paragraph(f"• {bullet.strip()}", body))

    if data.get("education"):
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in data["education"]:
            score = f" — Score: {_safe(edu.get('score'))}" if edu.get("score") else ""
            story.append(Paragraph(f"<b>{_safe(edu.get('degree'))} — {_safe(edu.get('institution'))}</b> ({_safe(edu.get('duration'))}){score}", body))

    if data.get("certifications"):
        story.append(Paragraph("CERTIFICATIONS", heading_style))
        for c in data["certifications"]:
            if c.strip():
                story.append(Paragraph(f"• {c.strip()}", body))

    doc.build(story)
    return str(out_path)
